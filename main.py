from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import latex2mathml.converter

def latex_to_omml(latex):
    mathml = latex2mathml.converter.convert(latex)
    omml = f'''
    <m:oMathPara xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
      <m:oMath>
        {mathml}
      </m:oMath>
    </m:oMathPara>
    '''
    return omml

def add_math(paragraph, equation):
    # Create the <m:oMathPara> element
    omath_para = OxmlElement('m:oMathPara')
    omath = OxmlElement('m:oMath')
    omath_para.append(omath)

    # Create the <m:r> element
    run = OxmlElement('m:r')
    run_properties = OxmlElement('m:rPr')
    run.append(run_properties)

    # Create the <m:t> element
    text = OxmlElement('m:t')
    text.text = equation
    run.append(text)

    # Append the run to the <m:oMath> element
    omath.append(run)

    # Append the <m:oMathPara> element to the paragraph
    paragraph._element.append(omath_para)

# Create a new Document
doc = Document()

# Add a title to the document
doc.add_heading('Math Equation Example', level=1)

# Add some text
p = doc.add_paragraph('Below is an example of a math equation:\n')

latex = r"E = mc^2"
omml = latex_to_omml(latex)
# Add a math equation
add_math(p, omml)

# Save the document
output_path = 'example.docx'
doc.save(output_path)

print(f"The Word document has been saved as {output_path}")
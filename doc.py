import latex2mathml.converter
from lxml import etree
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def latex_to_mathml(latex_str):
    return latex2mathml.converter.convert(latex_str)

def mathml_to_omml(mathml_str):
    xslt = etree.XML('''
    <xsl:stylesheet xmlns:xsl="http://www.w3.org/1999/XSL/Transform" version="1.0"
                    xmlns:m="http://www.w3.org/1998/Math/MathML"
                    xmlns:mc="http://schemas.openxmlformats.org/officeDocument/2006/math">
      <xsl:template match="/">
        <mc:mathPara>
          <mc:math>
            <xsl:apply-templates select="*"/>
          </mc:math>
        </mc:mathPara>
      </xsl:template>
      <xsl:template match="*">
        <xsl:element name="mc:{local-name()}">
          <xsl:apply-templates select="@* | node()"/>
        </xsl:element>
      </xsl:template>
      <xsl:template match="@*">
        <xsl:attribute name="mc:{local-name()}">
          <xsl:value-of select="."/>
        </xsl:attribute>
      </xsl:template>
    </xsl:stylesheet>
    ''')
    transform = etree.XSLT(xslt)
    mathml = etree.fromstring(mathml_str)
    omml = transform(mathml)
    return etree.tostring(omml)

def insert_omml_into_word(omml_str, output_file):
    document = Document()
    p = document.add_paragraph()
    r = p.add_run()
    r._r.append(parse_xml(omml_str))
    document.save(output_file)

latex_str = r'\frac{-b\pm\sqrt{b^{2}-4ac}}{2a}'
latex_str = r'\frac{a}{b}'

mathml_str = latex_to_mathml(latex_str)
omml_str = mathml_to_omml(mathml_str)
omml_str = omml_str.decode('utf-8')
insert_omml_into_word(omml_str, 'output.docx')